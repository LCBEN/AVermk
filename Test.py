from openpyxl import load_workbook
from docx import Document

# Excel-Tabelle öffnen
workbook = load_workbook('deine_tabelle.xlsx')
sheet = workbook.active

# Beispielhafte Verarbeitung der Werte
wert1 = sheet['A1'].value
wert2 = sheet['B1'].value

# Word-Dokument erstellen
document = Document()

# Text mit Formatierung hinzufügen
paragraph = document.add_paragraph()
run = paragraph.add_run()
run.text = f"Der Wert 1 ist {wert1} und der Wert 2 ist {wert2}."
run.bold = True
run.italic = True

# Word-Dokument speichern
document.save('ergebnis.docx')
